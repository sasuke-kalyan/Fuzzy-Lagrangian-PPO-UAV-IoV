"""Correct implementation-precision and formatting issues through methodology."""

from pathlib import Path

from docx import Document


SOURCE = Path(
    "/home/shubh-om/Downloads/projectdocnew/final one/"
    "thecodexone_report_formatted.docx"
)
OUTPUT = Path("thecodexone_report_methodology_corrected.docx")


def set_text(paragraph, text):
    """Replace paragraph text while retaining its paragraph-level formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def replace_text(document, old, new):
    count = 0
    for paragraph in all_paragraphs(document):
        if old not in paragraph.text:
            continue
        revised = paragraph.text.replace(old, new)
        set_text(paragraph, revised)
        count += 1
    return count


def append_to_matching_paragraph(document, prefix, addition):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            if addition not in paragraph.text:
                set_text(paragraph, paragraph.text.rstrip() + " " + addition)
            return True
    return False


def main():
    document = Document(SOURCE)

    # Abstract: describe the actual simulation design more precisely.
    replace_text(
        document,
        "uses synthetic urban-canyon, suburban-crossroads, and "
        "emergency-response datasets",
        "uses synthetically configured urban-canyon, suburban-crossroads, and "
        "emergency-response simulation environments",
    )

    # Introduction: make the organization statement grammatical and forward-looking.
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("The remainder of this paper is organized as follows."):
            set_text(
                paragraph,
                "The remainder of this paper is organized as follows. Section 2 reviews "
                "related work on UAV-assisted IoV communication, intelligent UAV "
                "selection, fuzzy decision systems, and constrained reinforcement "
                "learning. Section 3 presents the proposed Fuzzy LP-PPO methodology, "
                "including the system model, candidate screening, fuzzy-risk ranking, "
                "adaptive Lagrangian formulation, and PPO-based policy optimization. "
                "The subsequent sections will present the experimental setup, results "
                "and discussion, and conclusions.",
            )
            break

    # Avoid a premature results conclusion in the introduction.
    replace_text(
        document,
        "The experimental results demonstrate that the proposed framework achieves "
        "improved communication performance while effectively handling dynamic "
        "network conditions and multiple QoS constraints.",
        "The framework is evaluated using reward, delay, PDR, and signal-quality "
        "measurements against representative UAV-selection baselines.",
    )

    # Explain the precise observation contents and what remains outside the PPO state.
    append_to_matching_paragraph(
        document,
        "The processed Fuzzy LP-PPO observation",
        "The resulting 34-dimensional vector contains two normalized vehicle-position "
        "features and four features for each of eight UAVs: relative x-position, "
        "relative y-position, signal quality, and normalized residual energy. Delay, "
        "PDR, fuzzy priority, and Lagrangian multipliers are used in candidate "
        "evaluation or reward calculation but are not direct elements of this vector.",
    )

    # Clarify decision scope without changing the multi-vehicle scenario definition.
    append_to_matching_paragraph(
        document,
        "The input data required for the proposed Fuzzy LP-PPO methodology",
        "Although each scenario represents a multi-vehicle traffic environment, one "
        "representative active vehicle is served during each UAV-selection decision.",
    )

    # Align the scenario prose with the implemented mobility ranges and decision scope.
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("At the beginning of every episode"):
            set_text(
                paragraph,
                "At the beginning of every episode, the representative active vehicle "
                "is initialized randomly inside the operational area. During each "
                "communication decision, its position is updated by a scenario-specific "
                "random displacement. The next position is computed as in Equation (4), "
                "where Δp_v(t) denotes the displacement during the current decision, L "
                "is the side length of the simulation region, and clip(·) confines the "
                "updated coordinates to the operational boundary.",
            )
        elif paragraph.text.startswith("Urban-canyon scenario:"):
            set_text(
                paragraph,
                "Urban-canyon scenario: This scenario represents a dense IoV setting "
                "with approximately 20 vehicles, high traffic load, and comparatively "
                "large predominantly horizontal mobility increments. Its delay and PDR "
                "settings represent demanding communication conditions and emphasize "
                "collision-avoidance communication reliability.",
            )
        elif paragraph.text.startswith("Suburban-crossroads scenario:"):
            set_text(
                paragraph,
                "Suburban-crossroads scenario: This scenario represents a sparser "
                "15-vehicle road layout with moderate two-dimensional mobility. Its "
                "communication configuration applies the largest delay multiplier and "
                "energy-drain allowance, emphasizing endurance and energy-efficient UAV "
                "selection.",
            )
        elif paragraph.text.startswith("Emergency-response scenario:"):
            set_text(
                paragraph,
                "Emergency-response scenario: This scenario represents a dynamic "
                "15-vehicle incident-response setting with two-dimensional mobility and "
                "an emergency phase beginning approximately one-third of the way through "
                "the episode. It emphasizes adaptation to changing communication demand; "
                "the fuzzy emergency boost is activated through the scenario focus.",
            )

    # Clarify the mask, ranked action semantics, and empty-set fallback.
    append_to_matching_paragraph(
        document,
        "Each UAV is checked using the hard-feasibility condition",
        "The hard screen uses distance, normalized energy, signal, and delay; PDR "
        "remains a soft Lagrangian constraint. A PPO action denotes a rank in the "
        "current retained list, not a permanently assigned physical UAV, and is mapped "
        "to the corresponding UAV after screening and sorting.",
    )

    # Clarify energy timing and measurement stochasticity.
    append_to_matching_paragraph(
        document,
        "UAV residual energy evolves according to Equation (9)",
        "Reward and constraint calculations use the selected UAV's pre-drain energy "
        "E_u(t), whereas the logged remaining-energy value corresponds to the "
        "post-drain state E_u(t+1). Signal is distance based, while delay includes a "
        "small stochastic component; consequently, repeated link previews can vary "
        "slightly even when positions are unchanged.",
    )

    # Describe the implemented one-sided, persistent multiplier update exactly.
    append_to_matching_paragraph(
        document,
        "Four active soft QoS constraints are monitored",
        "The multipliers increase when their corresponding violations occur, remain "
        "unchanged when the constraints are satisfied, and are clipped to [0, 10].",
    )

    # Clarify scenario-specific training, sampling, deterministic inference and rollout.
    append_to_matching_paragraph(
        document,
        "The proposed model leverages the state representation",
        "A separate PPO policy is trained for each scenario. During training, actions "
        "are sampled from the policy distribution; during online evaluation, the "
        "highest-probability ranked action is selected.",
    )
    append_to_matching_paragraph(
        document,
        "Before ordinary PPO learning begins",
        "The configured maximum rollout length is 512 steps, while the implemented "
        "scenario rollout is min(512, 4 × 50) = 200 steps.",
    )

    # The emergency boost is scenario-level in the implementation.
    append_to_matching_paragraph(
        document,
        "Equations (17)–(20) define four normalized risk components",
        "In the implementation, the emergency boost is activated by the Emergency "
        "Response scenario's dynamic-adaptation focus; it is not derived from an "
        "independent per-step emergency flag.",
    )

    # Repair the malformed Equation 39 label without touching its Office Math content.
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "(39":
            set_text(paragraph, "(39)")

    # The final parameters table follows Table 5, so its correct label is Table 6.
    # Earlier architecture caption was accidentally labelled Table 6 and is Table 4.
    labels = [
        p for p in document.paragraphs
        if p.text.strip().startswith("Table ") and p.text.strip()[6:].isdigit()
    ]
    expected = ["Table 1", "Table 2", "Table 3", "Table 4", "Table 5", "Table 6"]
    for paragraph, label in zip(labels, expected):
        set_text(paragraph, label)

    # Remove empty Heading paragraphs that pollute a generated table of contents.
    for paragraph in list(document.paragraphs):
        if not paragraph.text.strip() and paragraph.style.name.startswith("Heading"):
            paragraph.style = document.styles["Normal"]

    # Normalize any remaining legacy method abbreviations.
    replacements = {
        "FLP-PPO": "Fuzzy LP-PPO",
        "FL-PPO": "Fuzzy LP-PPO",
        "fuzzy inference mechanism": "weighted clipped fuzzy-risk mechanism",
        "fuzzy inference": "clipped fuzzy-risk evaluation",
        "fuzzy membership mechanism": "weighted clipped fuzzy-risk mechanism",
    }
    for old, new in replacements.items():
        replace_text(document, old, new)

    document.core_properties.title = (
        "Fuzzy LP-PPO for Reliable UAV Selection in Dynamic UAV-Assisted IoV"
    )
    document.core_properties.comments = (
        "Corrected for implementation precision and formatting through the Proposed "
        "Methodology section. Results and conclusion sections were not added."
    )
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
