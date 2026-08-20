from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document


TARGET = Path("Fuzzy_LP_PPO_2 (1).docx")
SOURCE = Path("Fuzzy_LP_PPO_final.docx")
BACKUP = Path("Fuzzy_LP_PPO_2 (1)_before_table_update.docx")
DELIVERABLE = Path("Fuzzy_LP_PPO_2(1).docx")


def replace_table(target_table, source_table):
    target_table._element.getparent().replace(
        target_table._element, deepcopy(source_table._element)
    )


def paragraph_with_exact_text(document, text):
    return next(p for p in document.paragraphs if p.text.strip() == text)


def paragraph_starting(document, prefix):
    return next(p for p in document.paragraphs if p.text.strip().startswith(prefix))


def insert_after(element, new_element):
    element.addnext(deepcopy(new_element))


def delete_paragraphs_between(document, start_text, end_prefix):
    start = paragraph_with_exact_text(document, start_text)
    end = paragraph_starting(document, end_prefix)
    current = start._element.getnext()
    while current is not None and current is not end._element:
        following = current.getnext()
        current.getparent().remove(current)
        current = following


def main():
    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)

    target = Document(TARGET)
    source = Document(SOURCE)

    if len(target.tables) != 7 or len(source.tables) != 11:
        raise RuntimeError(
            f"Unexpected table counts: target={len(target.tables)}, source={len(source.tables)}"
        )

    # Replace the seven table objects already present in the preferred document.
    # Target semantics: Table 1, Table 5, Table 6, Table 7, Algorithms 2–3, Table 8.
    source_indices = [0, 4, 5, 6, 8, 9, 10]
    for target_table, source_index in zip(list(target.tables), source_indices):
        replace_table(target_table, source.tables[source_index])

    # Insert literature Tables 2 and 3 after the literature-review synthesis and
    # before the research-gap list. Copy their captions from the final document.
    literature_anchor = paragraph_starting(
        target,
        "The reviewed tells the effectiveness of Deep Reinforcement Learning",
    )
    caption_2 = paragraph_with_exact_text(
        source, "Table 2 Detailed review of 15 verified representative papers"
    )
    caption_3 = paragraph_with_exact_text(
        source, "Table 3 Comparison of related UAV-IoV decision and learning methods"
    )
    cursor = literature_anchor._element
    for element in (
        caption_2._element,
        source.tables[1]._element,
        caption_3._element,
        source.tables[2]._element,
    ):
        insert_after(cursor, element)
        cursor = cursor.getnext()

    # Insert preprocessing Table 4 immediately after the preprocessing heading.
    preprocessing_anchor = paragraph_with_exact_text(target, "Data pre-processing")
    caption_4 = paragraph_with_exact_text(
        source, "Table 4 Acquired data and preprocessing operations"
    )
    insert_after(preprocessing_anchor._element, caption_4._element)
    insert_after(preprocessing_anchor._element.getnext(), source.tables[3]._element)

    # Replace the paragraph-form Algorithm 1 with the formatted single-cell table
    # used in the final document, retaining the existing algorithm heading.
    delete_paragraphs_between(
        target,
        "Algorithm 1: Fuzzy-Based Candidate Screening and Ranking",
        "The algorithm performs a single traversal of all candidate UAVs",
    )
    algorithm_1_anchor = paragraph_with_exact_text(
        target, "Algorithm 1: Fuzzy-Based Candidate Screening and Ranking"
    )
    insert_after(algorithm_1_anchor._element, source.tables[7]._element)

    target.save(TARGET)
    shutil.copy2(TARGET, DELIVERABLE)
    print(TARGET)
    print(DELIVERABLE)


if __name__ == "__main__":
    main()
