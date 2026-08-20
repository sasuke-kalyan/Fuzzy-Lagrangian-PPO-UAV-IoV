"""Create a concise, implementation-faithful version of the uploaded report."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path("/home/shubh-om/Downloads/projectdocnew/2 (1).docx")
OUTPUT = Path("Fuzzy_LP_PPO_concise_project_report.docx")


def set_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def delete_table(table):
    table._element.getparent().remove(table._element)


def main():
    document = Document(SOURCE)
    original = list(document.paragraphs)

    # Mandatory title and a 100–250-word abstract grounded in the current code,
    # training logs, and common-evaluation results.
    set_text(
        original[1],
        "Fuzzy Lagrangian Proximal Policy Optimization for Reliable UAV Selection "
        "in Dynamic UAV-Assisted Internet of Vehicles",
    )
    set_text(
        original[5],
        "Dynamic vehicle mobility, time-varying links, and limited aerial energy make "
        "reliable unmanned aerial vehicle (UAV) selection difficult in the Internet of "
        "Vehicles (IoV). Fixed nearest-link, strongest-signal, and energy-greedy rules "
        "optimize only part of the quality-of-service (QoS) objective, while unconstrained "
        "reinforcement learning does not adapt the cost of repeated violations. This work "
        "proposes Fuzzy Lagrangian Proximal Policy Optimization (Fuzzy LP-PPO) for "
        "constraint-aware UAV selection. For each decision, the method acquires vehicle "
        "and eight-UAV positions, link signal, delay, packet delivery ratio (PDR), and "
        "residual energy. It normalizes the 34-dimensional observation, screens infeasible "
        "links, ranks retained candidates with clipped fuzzy-risk memberships, updates "
        "persistent Lagrangian multipliers, and learns a ranked action using PPO. The "
        "Python implementation uses synthetic urban-canyon, suburban-crossroads, and "
        "emergency-response datasets and trains one policy per scenario for 100 episodes "
        "of 50 steps. Under a common 20-episode evaluation, Fuzzy LP-PPO obtains a mean "
        "episode reward of 5415.31 versus 706.01 for the reward-aware heuristic, a 667.0% "
        "increase. It also reduces average delay from 81.81 to 75.78 ms (7.4%), increases "
        "PDR by 0.39 percentage points, and improves mean signal quality by 12.2%. The "
        "results demonstrate adaptive multi-QoS UAV selection under heterogeneous, "
        "non-stationary IoV conditions.",
    )

    # Make the end-of-introduction objectives explicit, as required.
    set_text(original[18], "The objectives of this work are:")
    set_text(
        original[19],
        "1. Develop Fuzzy LP-PPO to combine feasibility screening, clipped fuzzy-risk "
        "ranking, persistent Lagrangian penalties, and PPO for reliable UAV selection.",
    )
    set_text(
        original[20],
        "2. Model urban-canyon, suburban-crossroads, and emergency-response conditions "
        "with eight candidate UAVs and reproducible scenario-specific training.",
    )
    set_text(
        original[21],
        "3. Evaluate reward, delay, PDR, and signal quality against fixed UAV-selection "
        "baselines under identical environment dynamics and evaluation seeds.",
    )

    # Retain all 20 required paper descriptions, but remove two optional tables that
    # duplicate the narrative review. This is the main page-count reduction.
    delete_table(document.tables[2])
    delete_table(document.tables[1])
    for index in (49, 50, 51, 52):
        delete_paragraph(original[index])
    set_text(
        original[27],
        "This section reviews 20 recent related studies on UAV-assisted IoV, vehicular "
        "edge computing, and learning-based offloading. Each study is summarized by its "
        "method, environment, reported attainment, merit, and limitation before the "
        "research gaps are identified.",
    )
    set_text(
        original[48],
        "The 20 studies show that DRL, MARL, PPO, graph learning, and actor–critic "
        "methods can reduce delay and energy or improve throughput in simulated "
        "UAV-assisted vehicular systems. However, most optimize offloading, caching, "
        "trajectory, or resource allocation rather than reliable UAV selection, and few "
        "jointly represent uncertain link risk and adapt multiple QoS penalties. These "
        "limitations motivate the integrated Fuzzy LP-PPO method.",
    )

    # Give the proposed method a specific, reusable name and a clear section hierarchy.
    set_text(
        original[65],
        "Proposed Fuzzy Lagrangian PPO Methodology for Reliable UAV Selection (Fuzzy LP-PPO)",
    )
    original[65].style = document.styles["Heading 1"]
    methodology_headings = {
        68: "Data acquisition",
        70: "Data preprocessing",
        81: "System architecture and problem formulation",
        99: "Vehicle mobility and scenario model",
        113: "Communication, energy, and constraint model",
        141: "Fuzzy-priority estimation and candidate ranking",
        166: "Reward and adaptive Lagrangian formulation",
        175: "Policy optimization with PPO",
    }
    for index, text in methodology_headings.items():
        set_text(original[index], text)
        original[index].style = document.styles["Heading 2"]
    set_text(
        original[67],
        "The proposed Fuzzy LP-PPO methodology selects a reliable UAV in a dynamic "
        "UAV-IoV network. It combines clipped fuzzy-risk evaluation, PPO policy learning, "
        "and persistent Lagrangian penalties for delay, PDR, residual energy, and signal "
        "constraints. Figure 2 presents the topology, Figure 3 traces candidate screening "
        "and ranked action mapping, and Figure 4 shows the end-to-end network workflow.",
    )
    # Align the prose and Office Math thresholds with proposed/constraints.py.
    set_text(
        original[126],
        "Four active soft QoS constraints are monitored: delay ≤ 95 ms, PDR ≥ 55%, "
        "normalized residual energy ≥ 20%, and signal quality ≥ 0.10. Equation (11) "
        "defines non-negative violation magnitudes. Their multipliers persist across "
        "episode resets and are updated after every decision, so repeated violations "
        "receive progressively larger training penalties.",
    )
    for element in original[127]._element.iter(qn("m:t")):
        if not element.text:
            continue
        for old, new in {
            "-100)": "-95)",
            "48-": "55-",
            "8-": "20-",
            "0.05-": "0.10-",
        }.items():
            element.text = element.text.replace(old, new)

    # Condense prose that repeats the algorithms already printed in full.
    set_text(
        original[159],
        "Algorithm 1 converts instantaneous vehicle–UAV measurements into a feasible, "
        "ordered candidate list. Its inputs are vehicle/UAV positions, residual energies, "
        "and the emergency indicator; its output is the retained list C*. For every UAV "
        "it computes link metrics, the four risks in Equations (17)–(20), fuzzy priority "
        "in Equation (21), feasibility in Equation (22), and score in Equation (23). It "
        "then filters, sorts, and maps the PPO rank through Equation (24). If all links "
        "fail screening, the full fleet is retained to prevent an empty action set.",
    )
    set_text(
        original[160],
        "The ranking gives PPO a stable ordinal action space while remaining responsive "
        "to the current physical UAV set. Candidate evaluation requires O(N_U log N_U) "
        "time because sorting dominates, and O(N_U) auxiliary space.",
    )
    for index in range(161, 166):
        delete_paragraph(original[index])

    set_text(
        original[190],
        "Algorithm 2 trains one constraint-aware policy per scenario. It initializes the "
        "actor, critic, and persistent multiplier vector; performs the fuzzy-guided warm "
        "start; collects on-policy transitions; updates violations and multipliers at "
        "every step; computes GAE after each 200-step rollout; and applies ten clipped-PPO "
        "epochs. Its inputs and outputs are defined in the algorithm, and Equations "
        "(25)–(39) specify reward, dual ascent, advantage estimation, and optimization.",
    )
    set_text(
        original[191],
        "With N_ep episodes, horizon H, N_U candidates, PPO epochs E_PPO, and W "
        "actor–critic parameters, training requires O(E_WS N_WS W + N_ep H "
        "[N_U log N_U + E_PPO W]) time. Rollout, model, and candidate storage require "
        "O(K(d_s + 1) + W + N_U) space.",
    )
    for index in range(192, 198):
        delete_paragraph(original[index])

    set_text(
        original[200],
        "Algorithm 3 performs online selection with the trained scenario-specific actor. "
        "It normalizes the current measurements using Equations (12)–(16), obtains C* "
        "from Algorithm 1, selects the highest-probability rank, maps that rank to a UAV "
        "through Equation (24), and records the resulting QoS. Its decision-time cost is "
        "O(N_U log N_U + W_π), with O(N_U + d_s + h_max) space.",
    )
    for index in range(201, 207):
        delete_paragraph(original[index])

    # Remove the second, fully repetitive complexity subsection and equations.
    for index in range(208, 215):
        delete_paragraph(original[index])

    # Renumber retained tables after removing the two optional literature tables.
    table_labels = {
        71: "Table 2",
        129: "Table 3",
        133: "Table 4",
        151: "Table 5",
        215: "Table 6",
    }
    for index, label in table_labels.items():
        set_text(original[index], label)

    # Clean up terminology that could imply a Mamdani rule base not present in code.
    replacements = {
        "fuzzy inference mechanism": "clipped fuzzy-risk mechanism",
        "fuzzy inference": "clipped fuzzy-risk evaluation",
        "Fuzzy Logic-Based Lagrangian Proximal Policy Optimization (FLP-PPO)": (
            "Fuzzy Lagrangian Proximal Policy Optimization (Fuzzy LP-PPO)"
        ),
        "FL-PPO": "Fuzzy LP-PPO",
    }
    containers = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for paragraph in containers:
        text = paragraph.text
        revised = text
        for old, new in replacements.items():
            revised = revised.replace(old, new)
        if revised != text:
            set_text(paragraph, revised)

    document.core_properties.title = (
        "Concise Fuzzy LP-PPO Project Report for Reliable UAV Selection"
    )
    document.core_properties.comments = (
        "Shortened against NIT Raipur report-preparation instructions; "
        "implementation claims checked against the project workspace."
    )
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
