"""Apply only the user-requested metric and active-constraint corrections."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path("/home/shubh-om/Downloads/projectdocnew/thecodexone.docx")
OUTPUT = Path("thecodexone_corrected.docx")


def set_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main():
    document = Document(SOURCE)

    abstract = document.paragraphs[5]
    abstract_text = abstract.text
    replacements = {
        "7.8% reduction in average delay": "7.4% reduction in average delay",
        "14.8% improvement in signal quality": "12.2% improvement in signal quality",
    }
    for old, new in replacements.items():
        abstract_text = abstract_text.replace(old, new)
    set_text(abstract, abstract_text)

    constraint_paragraph = document.paragraphs[122]
    set_text(
        constraint_paragraph,
        "Four active soft QoS constraints are monitored: delay ≤ 95 ms, PDR ≥ 55%, "
        "normalized residual energy ≥ 20%, and signal quality ≥ 0.10. Equation (11) "
        "defines non-negative violation magnitudes. Their multipliers persist across "
        "episode resets and are updated after every decision, so repeated violations "
        "receive progressively larger training penalties.",
    )

    equation_11 = document.paragraphs[123]
    for element in equation_11._element.iter(qn("m:t")):
        if not element.text:
            continue
        for old, new in {
            "-100)": "-95)",
            "48-": "55-",
            "8-": "20-",
            "0.05-": "0.10-",
        }.items():
            element.text = element.text.replace(old, new)

    document.core_properties.comments = (
        "Only the requested delay, signal-improvement, and active QoS-constraint "
        "values were confirmed/corrected. Diagrams and all other content are unchanged."
    )
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
