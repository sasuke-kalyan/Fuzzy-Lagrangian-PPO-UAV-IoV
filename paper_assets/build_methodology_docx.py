#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

ROOT=Path(__file__).parent
OUT=ROOT/"proposed_methodology_package"
FIG=OUT/"figures"
OUT.mkdir(exist_ok=True)
d=Document(); s=d.sections[0]; s.top_margin=Inches(.8); s.bottom_margin=Inches(.9); s.left_margin=Inches(.8); s.right_margin=Inches(.8)
style=d.styles["Normal"]; style.font.name="Times New Roman"; style.font.size=Pt(11)

def p(t="",bold=False,italic=False,align=WD_ALIGN_PARAGRAPH.JUSTIFY,after=6):
    x=d.add_paragraph(); x.alignment=align; x.paragraph_format.line_spacing=1.15; x.paragraph_format.space_after=Pt(after)
    r=x.add_run(t); r.bold=bold; r.italic=italic; r.font.name="Times New Roman"; r.font.size=Pt(11); return x
def h(t,level=1):
    x=p(t,bold=True,align=WD_ALIGN_PARAGRAPH.LEFT,after=5); x.paragraph_format.space_before=Pt(10 if level==1 else 7); x.paragraph_format.keep_with_next=True; return x
def eq(t,n):
    x=p(f"{t}    ({n})",align=WD_ALIGN_PARAGRAPH.CENTER,after=7); x.runs[0].font.name="Cambria Math"; return x
def table(headers,rows,caption):
    p(caption,bold=True,align=WD_ALIGN_PARAGRAPH.LEFT,after=3)
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,v in enumerate(headers): t.rows[0].cells[i].text=v
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    for ri,row in enumerate(t.rows):
        for c in row.cells:
            for par in c.paragraphs:
                par.paragraph_format.space_after=Pt(1); par.paragraph_format.line_spacing=1
                for r in par.runs: r.font.name="Times New Roman"; r.font.size=Pt(8); r.bold=(ri==0)
    p("",after=2)
def fig(name,caption,note):
    p(f"PLACEMENT NOTE (remove after pasting): {note}",italic=True,align=WD_ALIGN_PARAGRAPH.LEFT,after=3)
    x=d.add_paragraph(); x.alignment=WD_ALIGN_PARAGRAPH.CENTER; x.add_run().add_picture(str(FIG/name),width=Inches(6.3))
    p(caption,italic=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
def alg(title,body):
    p(title,bold=True,align=WD_ALIGN_PARAGRAPH.LEFT,after=3)
    x=p(body,align=WD_ALIGN_PARAGRAPH.LEFT,after=8); x.paragraph_format.left_indent=Inches(.25)
    for r in x.runs: r.font.name="Courier New"; r.font.size=Pt(9)

h("5. PROPOSED METHODOLOGY")
h("5.1 Name of the Proposed Methodology",2)
p("This work proposes the Fuzzy Lagrangian Proximal Policy Optimization Methodology for Reliable UAV Selection in Dynamic Internet of Vehicles, abbreviated as Fuzzy LP-PPO. The short name “Fuzzy LP-PPO” is used in algorithms, tables, graphs, and performance comparisons.")
p("The methodology selects a suitable UAV for a moving vehicle by jointly considering vehicle–UAV distance, communication delay, packet delivery ratio (PDR), signal quality, and residual UAV energy. It combines fuzzy logic for gradual risk and urgency estimation, hard feasibility screening, adaptive Lagrangian multipliers, stable PPO learning, and scenario-specific training for urban-canyon, suburban-crossroads, and emergency-response environments.")

h("5.2 Data Acquisition and Scenario Generation",2)
p("The proposed system uses a simulation-based data-acquisition process. At the beginning of every episode, the vehicle position, UAV positions, UAV altitude, and residual UAV energy are initialized using a seeded random generator. The operational region is a 2000 × 2000 m² area containing eight candidate UAVs. UAV altitude is sampled between 80 and 200 m, while initial UAV energy is sampled between 30 and 50 internal units.")
p("The urban-canyon scenario represents dense corridor movement and obstruction-sensitive links; the suburban-crossroads scenario represents sparse movement, longer links, and energy-aware operation; and the emergency-response scenario represents irregular mobility and non-stationary demand. At each step, the system acquires vehicle coordinates, UAV coordinates, relative position, distance, signal, delay, PDR, energy, scenario type, and current Lagrangian multipliers.")
table(["Input variable","Raw range/unit","Preprocessing","Purpose"],[
 ["Vehicle coordinates","0–2000 m","Divide by area size","Vehicle location"],["Relative UAV coordinates","−2000–2000 m","Clip and map to [0,1]","Relative geometry"],
 ["UAV altitude","80–200 m","Use in 3-D distance","Aerial geometry"],["Signal quality","0.05–1.00","Already normalized","Link strength"],
 ["Delay","0–100 ms","Threshold/fuzzy evaluation","Responsiveness"],["PDR","50–100%","Percentage form","Reliability"],
 ["UAV energy","0–50 units","Multiply by two and clip","Energy percentage"],["Lagrangian multipliers","0–10","Clip after update","Constraint importance"]],"Table 4. Acquired data and preprocessing operations")

h("5.3 Data Preprocessing",2)
p("Let L = 2000 m denote the side length of the operational region. The normalized vehicle coordinates are defined by Equation (1). Relative UAV coordinates are clipped and mapped to [0,1] through Equations (2) and (3), while Equation (4) converts internal UAV energy into a percentage-compatible value.")
eq("x̄ᵥ(t)=xᵥ(t)/L,     ȳᵥ(t)=yᵥ(t)/L",1)
eq("x̄ᵤ,ᵥ(t)=½[clip((xᵤ(t)−xᵥ(t))/L,−1,1)+1]",2)
eq("ȳᵤ,ᵥ(t)=½[clip((yᵤ(t)−yᵥ(t))/L,−1,1)+1]",3)
eq("Eᵤ%(t)=min(100,max(0,2Eᵤ(t)))",4)
p("The processed PPO observation is given by Equation (5). It contains two normalized vehicle features and four features for each of eight UAVs, producing a 34-dimensional observation.")
eq("sₜ=[x̄ᵥ,ȳᵥ,{x̄ᵤ,ᵥ,ȳᵤ,ᵥ,Sᵥ,ᵤ,Eᵤ/50}ᵤ₌₁⁸]",5)

h("5.4 Proposed System Architecture",2)
p("The complete Fuzzy LP-PPO architecture is presented in Figure 2. The communication module calculates link measurements, the hard mask removes clearly infeasible UAVs, the fuzzy module assigns continuous priority, the ranking module orders candidates, and PPO selects a ranked action. The reward and Lagrangian module then provides constraint-aware feedback.")
fig("figure_2_system_architecture.png","Figure 2. Proposed Fuzzy LP-PPO architecture for reliable UAV selection in dynamic IoV.","Insert Figure 2 immediately after the first architecture-description paragraph in Section 5.4.")
table(["Component","Input","Processing","Output"],[
 ["Scenario generator","Scenario ID and seed","Initialize mobility and UAVs","Environment state"],["Communication module","Coordinates","Distance, signal, delay, PDR","Link measurements"],
 ["Hard mask","Distance, delay, signal, energy","Apply thresholds","Feasible set"],["Fuzzy module","QoS, energy, scenario","Estimate gradual risk","Priority Fᵥ,ᵤ"],
 ["Ranking module","Measurements and priority","Composite scoring","Ranked UAV list"],["PPO actor–critic","State sₜ","Action probability/value","Action and Vφ(sₜ)"],
 ["Lagrangian module","Violations","Multiplier update","Adaptive penalty"],["Logging module","Action and metrics","Store records","Evaluation output"]],"Table 5. Components of the proposed architecture")

h("5.5 Fuzzy-Priority Estimation",2)
p("The fuzzy-priority module converts delay, PDR, signal, energy, and emergency information into a continuous risk-aware priority. Equations (6)–(9) define the four normalized risk components, and Equation (10) combines them.")
eq("μᴰ(t)=clip((Dᵥ,ᵤ(t)−55)/45,0,1)",6); eq("μᴾ(t)=clip((55−Pᵥ,ᵤ(t))/7,0,1)",7)
eq("μˢ(t)=clip((0.18−Sᵥ,ᵤ(t))/0.13,0,1)",8); eq("μᴱ(t)=clip((35−Eᵤ%(t))/35,0,1)",9)
eq("Fᵥ,ᵤ(t)=clip[0.35μᴰ+0.25μᴾ+0.25μˢ+0.15μᴱ+0.10Iₑₘg,0,1]",10)
table(["Fuzzy input","Low risk","Increasing risk","Maximum risk"],[["Delay","D≤55 ms","55–100 ms","D≥100 ms"],["PDR","P≥55%","48–55%","P≤48%"],["Signal","S≥0.18","0.05–0.18","S≤0.05"],["Energy","E%≥35%","0–35%","E%=0"],["Emergency","Normal","—","Emergency"]],"Table 6. Fuzzy-priority inputs")

h("5.6 Feasibility Screening and Candidate Ranking",2)
p("Each UAV is checked using Equation (11). If no UAV passes the conditions, the entire fleet is retained so that the action set never becomes empty. Feasible candidates are scored using Equation (12), and PPO selects a position in the ranked list as defined in Equation (13).")
eq("Mᵤ(t)=1 if dᵥ,ᵤ≤500, Eᵤ%≥20, Sᵥ,ᵤ≥0.10, and Dᵥ,ᵤ≤95; otherwise Mᵤ(t)=0",11)
eq("Cᵥ,ᵤ=240Sᵥ,ᵤ+2Pᵥ,ᵤ−1.3Dᵥ,ᵤ+2.5Eᵤ%+35Fᵥ,ᵤ−0.02dᵥ,ᵤ",12)
eq("aₜ ∈ {0,1,…,7}",13)
p("Figure 3 traces the processing performed for each UAV. The procedure is repeated for all eight candidates before sorting.")
fig("figure_3_candidate_trace.png","Figure 3. Trace-out diagram of communication measurement, fuzzy evaluation, and candidate ranking.","Insert Figure 3 after the feasibility and candidate-ranking equations in Section 5.6.")

h("5.7 Reward and Adaptive Lagrangian Formulation",2)
p("The QoS score, task-service reward, adaptive constraint penalty, and final reward are defined in Equations (14)–(17). Equation (18) gives the initial multipliers, while Equation (19) increases the multiplier associated with any observed violation.")
eq("Qᵥ,ᵤ(t)=190Sᵥ,ᵤ(t)+1.5Pᵥ,ᵤ(t)−1.2Dᵥ,ᵤ(t)+2.2Eᵤ%(t)",14)
eq("Rₜₐₛₖ(t)=15N꜀(t)−0.002dᵥ,ᵤ(t)−0.08eᵤ(t)−0.05",15)
eq("Φ(t)=λᴰgᴰ(t)+λᴾgᴾ(t)+λᴱgᴱ(t)+λˢgˢ(t)",16)
eq("rₜ=Rₜₐₛₖ(t)+Qᵥ,ᵤ(t)+20Fᵥ,ᵤ(t)−Φ(t)",17)
eq("λ(0)=[0.35,0.20,0.60,0.35]",18)
eq("λᵢ(t+1)=clip[λᵢ(t)+ηλgᵢ(t),0,10],     ηλ=0.03",19)

h("5.8 PPO Policy Learning and Warm Start",2)
p("The actor probability and critic value are represented by Equations (20) and (21). PPO uses the probability ratio, TD residual, generalized advantage estimate, clipped objective, critic loss, and complete loss in Equations (22)–(27).")
eq("πθ(aₜ|sₜ)",20); eq("Vφ(sₜ)=E[Σₖγᵏrₜ₊ₖ | sₜ]",21); eq("ρₜ(θ)=πθ(aₜ|sₜ)/πθₒₗd(aₜ|sₜ)",22)
eq("δₜ=rₜ+γVφ(sₜ₊₁)−Vφ(sₜ)",23); eq("Âₜ=Σₗ(γλGAE)ˡδₜ₊ₗ",24)
eq("Lclip(θ)=E[min(ρₜÂₜ,clip(ρₜ,1−ε,1+ε)Âₜ)]",25); eq("LV(φ)=E[(Vφ(sₜ)−R̂ₜ)²]",26)
eq("L(θ,φ)=−Lclip(θ)+cVLV(φ)−cH H[πθ(·|sₜ)]",27)
p("Before ordinary PPO learning, 800 states are sampled and the fuzzy-ranked best candidate is stored as action zero. Six behavior-cloning epochs minimize the warm-start loss in Equation (28), reducing unstructured initial exploration.")
eq("LWS(θ)=−(1/NWS)Σⱼ log πθ(aⱼF|sⱼ),     aⱼF=0",28)

h("5.9 Training Flowchart and Sequence of Operation",2)
p("Figure 4 presents the complete training flow. Candidate evaluation and multiplier adaptation occur at every step, whereas PPO parameters are updated after rollout collection.")
fig("figure_4_training_flowchart.png","Figure 4. Training flowchart of the proposed Fuzzy LP-PPO methodology.","Insert Figure 4 after the PPO and warm-start descriptions in Section 5.9.")
p("Figure 5 presents the sequence of messages and computations during one vehicle-to-UAV selection decision.")
fig("figure_5_sequence_diagram.png","Figure 5. Sequence diagram of one Fuzzy LP-PPO UAV-selection decision.","Insert Figure 5 immediately after the sequence-operation paragraph in Section 5.9.")

h("5.10 Proposed Algorithms",2)
alg("Algorithm 1. Fuzzy-based candidate screening and ranking","Input: pᵥ(t), {pᵤ(t)}, {Eᵤ(t)}, Iₑₘg(t)\nOutput: Ranked candidate list Cₜ\n1: Cₜ ← ∅\n2: for each UAV u=1,…,8 do\n3:   Compute distance, signal, delay, PDR and normalized energy\n4:   Compute μᴰ, μᴾ, μˢ and μᴱ using Equations (6)–(9)\n5:   Compute Fᵥ,ᵤ using Equation (10)\n6:   Evaluate feasibility Mᵤ using Equation (11)\n7:   Compute Cᵥ,ᵤ using Equation (12) and store candidate\n8: end for\n9: If at least one candidate is feasible, remove infeasible candidates\n10: Sort by score, energy, signal and lower delay\n11: return Cₜ")
p("Algorithm 1 evaluates all eight UAVs and guarantees a non-empty candidate list through its fallback rule.")
alg("Algorithm 2. Adaptive Fuzzy LP-PPO training","Input: scenarios Ω, episodes Nₑₚ, horizon H, actor θ, critic φ, ηλ\nOutput: trained policies {πθ*}\n1: for each scenario ω∈Ω do\n2:   Initialize actor, critic and λ=[0.35,0.20,0.60,0.35]\n3:   Perform fuzzy-guided warm start\n4:   for episode=1,…,Nₑₚ do\n5:     Reset environment and rollout buffer\n6:     for t=0,…,H−1 do\n7:       Observe sₜ and execute Algorithm 1\n8:       Sample aₜ~πθ(aₜ|sₜ) and map it to the effective UAV\n9:       Execute communication and calculate rₜ\n10:      Compute violations and update λ using Equation (19)\n11:      Update energy, vehicle position and next state\n12:      Store transition in the rollout buffer\n13:    end for\n14:    Compute GAE and update actor and critic\n15:    Record episode metrics\n16:  end for\n17:  Save scenario policy\n18: end for")
p("Algorithm 2 describes the complete training process, including fuzzy warm start, Lagrangian adaptation, and PPO optimization.")
alg("Algorithm 3. Online UAV selection","Input: trained policy πθ* and current network state\nOutput: selected UAV u*\n1: Construct normalized observation sₜ\n2: Execute Algorithm 1 to obtain Cₜ\n3: Compute πθ*(·|sₜ)\n4: Select aₜ=arg maxₐ πθ*(a|sₜ)\n5: Map aₜ to effective UAV u*\n6: Establish communication and record QoS information\n7: return u*")

h("5.11 Computational Complexity and Implementation Parameters",2)
p("Communication and fuzzy evaluation require O(NU) operations, while ranking requires Equation (29). If Wπ denotes the number of actor parameters, the total deployment-time complexity is given by Equation (30).")
eq("O(NU log NU)",29); eq("O(NU log NU + Wπ)",30)
table(["Parameter","Value"],[["Candidate UAVs","8"],["Observation dimension","34"],["Actions","8"],["Actor/Critic","128–128 neurons"],["PPO learning rate","3×10⁻⁴"],["γ / λGAE / ε","0.99 / 0.95 / 0.20"],["Entropy coefficient","0.02"],["Batch size / PPO epochs","64 / 10"],["Warm-start samples / epochs","800 / 6"],["Episodes / steps","100 / 50 per scenario"],["Lagrangian step/range","0.03 / [0,10]"]],"Table 7. Principal implementation parameters")

h("5.12 Output of the Proposed Methodology",2)
p("For every decision, the proposed methodology produces the requested PPO action, effective selected UAV, ranked candidate list, distance, signal quality, delay, PDR, remaining UAV energy, immediate reward, violation magnitudes, updated multipliers, and episode-level reward. These outputs support reliability, energy, convergence, scenario, and policy-comparison analyses.")

d.core_properties.title="Proposed Fuzzy LP-PPO Methodology—Copy-Paste Research Paper Material"
d.save(OUT/"proposed_methodology_copy_paste.docx")
print(OUT/"proposed_methodology_copy_paste.docx")
