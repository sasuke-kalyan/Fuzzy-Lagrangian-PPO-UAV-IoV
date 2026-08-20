from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path("Fuzzy_LP_PPO_2.docx")
OUTPUT = Path("Fuzzy_LP_PPO_2_corrected.docx")


def replace_everywhere(document, old, new):
    """Replace text in ordinary paragraphs and table cells."""
    containers = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)

    count = 0
    for paragraph in containers:
        if old not in paragraph.text:
            continue
        # Most body paragraphs are single runs; this preserves their formatting.
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                count += 1
        # Handle terms split across runs while retaining the first-run formatting.
        if old in paragraph.text:
            text = paragraph.text.replace(old, new)
            if paragraph.runs:
                paragraph.runs[0].text = text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(text)
            count += 1
    return count


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_table_row(table, row_index):
    row = table.rows[row_index]
    row._element.getparent().remove(row._element)


def append_symbol_row(table, symbol, description):
    row = table.add_row()
    row.cells[0].text = symbol
    row.cells[1].text = description


def main():
    document = Document(SOURCE)

    # Report the actual experiment implemented by the project and recorded in logs.
    replace_everywhere(document, "1000 episodes of 1000 steps per scenario", "100 episodes of 50 steps per scenario")
    replace_everywhere(document, "1000 episodes, with 1000 interaction steps in every episode", "100 episodes, with 50 interaction steps in every episode")
    replace_everywhere(document, "1000 per scenario", "100 per scenario")
    replace_everywhere(document, "1000 / 1000", "100 / 50")

    # Correct table cross-references.
    replace_everywhere(document, "Table 3 lists 20 recent and representative studies", "Table 2 lists 16 representative studies")
    replace_everywhere(document, "Table 2 Detailed review of 20 recent and representative papers", "Table 2 Detailed review of 16 representative papers")
    replace_everywhere(document, "Table 9 summarizes", "Table 8 summarizes")

    # Align terminology with the implemented clipped linear memberships and weighted sum.
    terminology = {
        "fuzzy priority rules": "clipped fuzzy-risk priority mappings",
        "fuzzy rules estimate": "clipped linear risk memberships estimate",
        "fixed fuzzy rules": "fixed fuzzy mappings",
        "fuzzy inference system": "weighted clipped risk-priority module",
        "fuzzy inference mechanism": "weighted clipped risk-priority mechanism",
        "fuzzy linguistic reasoning": "clipped linear risk memberships",
        "fuzzy inference for interpreting uncertain communication conditions": "weighted clipped fuzzy-risk aggregation for representing gradual communication risk",
        "communication and fuzzy-priority evaluation": "communication and clipped risk-priority evaluation",
        "Communication and fuzzy evaluation": "Communication and clipped risk evaluation",
    }
    for old, new in terminology.items():
        replace_everywhere(document, old, new)

    # Table 1: clarify the decision scope and add the missing GAE coefficient.
    symbols = document.tables[0]
    symbols.cell(4, 1).text = "Set of active vehicles at decision step t; one representative active vehicle is served per selection decision"
    append_symbol_row(symbols, "λ_GAE", "Generalized advantage-estimation parameter")

    # Table 2: remove four unsupported rows and make the two retained legacy rows
    # faithful to the cited publications. Delete bottom-up to preserve indices.
    literature = document.tables[1]
    for index in (20, 19, 18, 15):
        delete_table_row(literature, index)

    # After deletion, Chen and Wang are rows 15 and 16.
    chen = literature.rows[15].cells
    chen[0].text = "Chen et al. 2016"
    chen[1].text = "Game-theoretic distributed multi-user computation offloading for mobile-edge cloud computing."
    chen[2].text = "Multi-user mobile-edge cloud computing simulation."
    chen[3].text = "Demonstrated efficient distributed offloading with low coordination overhead."
    chen[4].text = "Relevant optimization baseline for computation offloading."
    chen[5].text = "Not a deep-RL or direct UAV-selection method."

    wang = literature.rows[16].cells
    wang[0].text = "Wang et al. 2022"
    wang[1].text = "Delay-aware coordination of dependent microservices in mobile edge computing."
    wang[2].text = "Mobile edge-computing service-coordination environment."
    wang[3].text = "Reduced service delay through dependency-aware microservice coordination."
    wang[4].text = "Relevant delay-aware edge-service baseline."
    wang[5].text = "Not a multi-agent RL or UAV-IoV resource-allocation method."

    # Table 3: remove the unsupported named LC-MAPO comparison and describe the
    # implementation precisely rather than as a Mamdani/rule-base system.
    comparison = document.tables[2]
    delete_table_row(comparison, 10)
    comparison.rows[10].cells[0].text = "Clipped fuzzy-risk selection"
    comparison.rows[10].cells[1].text = "Weighted clipped linear risk memberships"
    comparison.rows[10].cells[2].text = "Interpretable and computationally efficient"
    comparison.rows[10].cells[3].text = "Fixed mappings do not learn long-term policies or adapt constraint prices"
    comparison.rows[11].cells[1].text = "Weighted clipped risk-priority aggregation, adaptive Lagrangian penalties, and PPO learning"

    # Table 4: distinguish broad formula bounds from values realized by a scenario.
    preprocessing = document.tables[3]
    preprocessing.rows[0].cells[1].text = "Model bound / unit"
    preprocessing.rows[4].cells[1].text = "0.05–1.00 (nominal bound)"
    preprocessing.rows[5].cells[1].text = "0–100 ms (capped model bound)"
    preprocessing.rows[6].cells[1].text = "50–100% (nominal bound)"
    preprocessing.rows[8].cells[2].text = "Internal state; clip after update"
    preprocessing.rows[8].cells[3].text = "Adaptive constraint importance; not part of the 34-dimensional PPO observation"

    # Table 5: actual run length and rollout size.
    system_parameters = document.tables[4]
    system_parameters.rows[10].cells[1].text = "100 per scenario"
    system_parameters.rows[11].cells[1].text = "50"
    rollout = system_parameters.add_row()
    rollout.cells[0].text = "PPO rollout length (n_steps)"
    rollout.cells[1].text = "200"

    # Table 6: exact module inputs and responsibilities.
    architecture = document.tables[5]
    architecture.rows[2].cells[1].text = "Vehicle/UAV coordinates"
    architecture.rows[2].cells[2].text = "Compute distance, signal, delay, and PDR"
    architecture.rows[3].cells[1].text = "Distance, delay, signal, and energy"
    architecture.rows[4].cells[0].text = "Weighted clipped risk-priority module"
    architecture.rows[4].cells[1].text = "Delay, PDR, signal, energy, and emergency indicator"
    architecture.rows[4].cells[2].text = "Compute clipped linear risks and their weighted sum"
    architecture.rows[4].cells[3].text = "Risk-aware priority F_v,u"

    # Table 8 (the last parameter table; tables 8–10 in the DOCX are algorithms).
    implementation = document.tables[10]
    implementation.rows[10].cells[1].text = "100 / 50"
    new_row = implementation.add_row()
    new_row.cells[0].text = "PPO rollout length (n_steps)"
    new_row.cells[1].text = "200"

    # Remove unsupported incomplete literature entries and rewrite dependent prose.
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("[15] ") or paragraph.text.startswith("[18] ") or paragraph.text.startswith("[19] ") or paragraph.text.startswith("[20] "):
            paragraph._element.getparent().remove(paragraph._element)

    replacements = {
        "(Chen et al. 2016; Wang et al. 2022; Zhang et al. 2022; Khan et al. 2023)": "(Chen et al. 2016; Wang et al. 2022)",
        "(Khan et al. 2023)": "",
        "(Zadeh 1965; Khan et al. 2023)": "(Zadeh 1965)",
        "(Khan et al. 2023).": ".",
        "(MADDPG, MAPPO/HAPPO, LC-MAPO)": "(MADDPG and MAPPO/HAPPO)",
        " (Khan et al. 2023)": "",
    }
    for old, new in replacements.items():
        replace_everywhere(document, old, new)

    # Replace the unsupported LC-MAPO-specific literature claim with a supported
    # constrained-RL summary.
    for paragraph in document.paragraphs:
        if "The LC-MAPO method applies this principle" in paragraph.text:
            set_paragraph_text(
                paragraph,
                "Safety and QoS constraints are critical when learned policies are deployed in transportation and aerial networks. Safe RL commonly models the problem as a constrained Markov decision process in which the policy maximizes expected reward while maintaining one or more expected costs below predefined limits (Altman 1999). Constrained policy optimization (CPO) restricts each policy update so that the new policy remains within an approximately feasible region (Achiam et al. 2017), while reward-constrained policy optimization (RCPO) introduces Lagrangian multipliers that adapt the weight assigned to constraint costs (Tessler et al. 2018). Lagrangian methods are attractive because they transform a constrained objective into a form that can be optimized using standard policy-gradient updates. The present work applies this principle to link-level delay, PDR, signal-quality, and residual-energy violations in a centralized UAV-selection problem."
            )

    # Correct overly strong linguistic-rule description in the fuzzy survey.
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("Fuzzy-logic systems are widely used"):
            set_paragraph_text(
                paragraph,
                "Fuzzy-logic concepts are useful when communication indicators are uncertain or do not have sharp boundaries (Zadeh 1965). Gradual membership values can represent transitions between safe and risky delay, signal, PDR, and energy conditions. In the present implementation, each risk is computed by a clipped linear membership function and the values are combined through a fixed weighted sum; no Mamdani rule base or centroid defuzzification is used. The resulting priority can guide initial decisions, shape the reward, and rank candidate actions, while reinforcement learning learns a policy from long-term interaction. Adaptive Lagrangian multipliers complement these fixed risk mappings by increasing the cost of repeated delay, PDR, signal, or energy violations during training."
            )

    document.core_properties.title = "Fuzzy LP-PPO for Reliable UAV Selection in Dynamic UAV-IoV Systems"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
