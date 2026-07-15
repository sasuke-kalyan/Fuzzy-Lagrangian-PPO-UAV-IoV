#!/usr/bin/env python3
"""Create a reference-styled DOCX containing only Sections 1–2 and references."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


FONT = "Droid Sans Mono"
FONT_SIZE = Pt(10.5)


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    text = re.sub(r"(?<!\\)\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    return text.strip()


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, bold=False, italic=False, size=FONT_SIZE) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = size
    run.bold = bold
    run.italic = italic


def add_paragraph(document, text="", *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  before=0, after=5, left=0, first=0, keep=False, size=FONT_SIZE):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(14.25)
    fmt.left_indent = Inches(left)
    fmt.first_line_indent = Inches(first)
    fmt.keep_with_next = keep
    run = paragraph.add_run(clean_inline(text))
    format_run(run, bold=bold, size=size)
    return paragraph


def add_heading(document, text, level):
    before = 10 if level == 1 else 7
    return add_paragraph(
        document,
        text,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        before=before,
        after=4,
        keep=True,
    )


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [part.strip() for part in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in row):
            rows.append(row)
        i += 1
    return rows, i


def add_table(document, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = clean_inline(row[c_idx] if c_idx < len(row) else "")
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = Pt(10)
            run = p.add_run(text)
            format_run(run, bold=(r_idx == 0), size=Pt(7.5))
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for edge in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{edge}")
                node.set(qn("w:w"), "55")
                node.set(qn("w:type"), "dxa")
                margins.append(node)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def selected_lines(source: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    sec3 = next(i for i, line in enumerate(lines) if line.startswith("## 3 "))
    refs = next(i for i, line in enumerate(lines) if line.strip() == "## References")
    return lines[:sec3] + [""] + lines[refs:]


def convert(source: Path, output: Path):
    lines = selected_lines(source)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.268)
    section.page_height = Inches(11.693)
    section.left_margin = Inches(0.7875)
    section.right_margin = Inches(0.7875)
    section.top_margin = Inches(0.7875)
    section.bottom_margin = Inches(1.1757)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = FONT_SIZE

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            i += 1
            continue

        image = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image:
            path = (source.parent / image.group(2)).resolve()
            if path.exists():
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(3)
                p.add_run().add_picture(str(path), width=Inches(6.0))
            i += 1
            continue

        if stripped == "$$":
            equation = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                equation.append(lines[i].strip())
                i += 1
            p = add_paragraph(document, "\n".join(equation), align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
            for run in p.runs:
                run.font.name = "Cambria Math"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1:
                add_paragraph(document, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                              after=9, keep=True, size=Pt(10.5))
            elif text == "Abstract":
                add_heading(document, text, 1)
            elif level == 2:
                add_heading(document, text, 1)
            else:
                add_heading(document, text, 2)
            i += 1
            continue

        if stripped.startswith("Extended author information"):
            i += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if ordered:
            add_paragraph(document, f"{ordered.group(1)}. {ordered.group(2)}",
                          left=0.25, first=-0.25, after=5)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            add_paragraph(document, f"• {bullet.group(1)}", left=0.25, first=-0.18, after=4)
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt == "$$" or nxt.startswith("#") or nxt.startswith("|")
                    or nxt.startswith("© ")
                    or re.fullmatch(r"!\[([^]]*)\]\([^)]+\)", nxt)
                    or re.match(r"^\d+\.\s+", nxt) or re.match(r"^[-*]\s+", nxt)):
                break
            paragraph_lines.append(nxt)
            i += 1
        text = " ".join(paragraph_lines)
        centered = text.startswith(("Shubh Om", "Received:", "© The Author"))
        add_paragraph(document, text, align=(WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY))

    document.core_properties.title = "Fuzzy LP-PPO UAV-IoV: Title, Abstract, Introduction and Related Work"
    document.core_properties.subject = "Documentation sections 1–4 only"
    document.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: create_styled_sections_docx.py SOURCE.md OUTPUT.docx")
    convert(Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve())
