import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path("Fuzzy_LP_PPO_2_corrected.docx")
OUTPUT = Path("Fuzzy_LP_PPO_final_implementation_faithful.docx")


def set_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_all(document, old, new):
    containers = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for paragraph in containers:
        if old not in paragraph.text:
            continue
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
        if old in paragraph.text:
            set_text(paragraph, paragraph.text.replace(old, new))


def delete_row(table, index):
    row = table.rows[index]
    row._element.getparent().remove(row._element)


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def paragraph_starting(document, prefix):
    return next(p for p in document.paragraphs if p.text.startswith(prefix))


def main():
    document = Document(SOURCE)

    # Newly regenerated common-evaluation values: PPO vs reward-aware selector.
    abstract = document.paragraphs[4]
    set_text(
        abstract,
        "The growth of connected vehicles and time-sensitive Internet of Vehicles (IoV) services has increased demand for reliable, low-latency communication. Unmanned aerial vehicle (UAV)-assisted networking can extend coverage, but selecting an appropriate UAV is difficult when mobility, link quality, and residual energy vary dynamically. Conventional greedy selection and unconstrained learning methods struggle to balance communication quality, energy availability, and quality-of-service (QoS) requirements. This paper addresses UAV selection while jointly controlling delay, packet delivery ratio (PDR), signal quality, and energy constraints. The proposed method, Fuzzy Lagrangian Proximal Policy Optimization (Fuzzy LP-PPO), combines clipped fuzzy-risk priority mappings, persistent adaptive Lagrangian penalties, and scenario-specific PPO learning. The framework prioritizes risky links, updates constraint multipliers across episode boundaries, and learns UAV-selection policies across eight candidate UAVs. Simulations cover urban-canyon, suburban-crossroads, and emergency-response environments, with policies trained for 100 episodes of 50 steps per scenario. Under a common 20-episode evaluation, the trained policy achieves a 667.0% increase in mean episode reward, a 7.4% reduction in average delay, a 0.39-percentage-point increase in PDR, and a 12.2% improvement in signal quality relative to the reward-aware heuristic. These results indicate that Fuzzy LP-PPO provides an effective basis for adaptive, constraint-aware communication in heterogeneous, non-stationary UAV-IoV environments."
    )

    results = paragraph_starting(document, "The trained model is compared under a common evaluation")
    set_text(
        results,
        "The trained model is compared under a common 20-episode evaluation with random, nearest-UAV, strongest-signal, energy-aware greedy, and reward-aware selection strategies. Averaged across the three scenarios, the trained policy obtains a mean episode reward of 5415.31 compared with 706.01 for the reward-aware heuristic, an improvement of approximately 667.0%. It reduces average delay from 81.81 ms to 75.78 ms (7.4%), increases PDR from 50.40% to 50.79% (0.39 percentage points), and raises mean signal quality from 0.1128 to 0.1266 (12.2%). The comparison uses identical environment dynamics, episode horizons, reward definitions, and evaluation seeds for all policies."
    )

    # Active soft constraints and persistence across episode boundaries.
    constraint_text = paragraph_starting(document, "To encourage long-term communication reliability")
    set_text(
        constraint_text,
        "To encourage long-term communication reliability, four active soft QoS constraints are monitored throughout reinforcement learning: maximum communication delay of 95 ms, minimum PDR of 55%, minimum normalized UAV energy of 20%, and minimum communication signal quality of 0.10. The non-negative violation functions in Equation (11) are zero when a requirement is satisfied and positive otherwise. Because these thresholds lie inside the attainable communication-model ranges, every constraint can provide a learning signal. The associated non-negative multipliers persist across episode resets and are updated after every communication decision, allowing repeated violations to accumulate progressively larger adaptive penalties throughout scenario training."
    )

    # Update the four constants embedded in the Office Math equation (11).
    equation_11 = document.paragraphs[116]
    math_replacements = {
        "-100)": "-95)",
        "48-": "55-",
        "8-": "20-",
        "0.05-": "0.10-",
    }
    for element in equation_11._element.iter(qn("m:t")):
        if not element.text:
            continue
        for old, new in math_replacements.items():
            element.text = element.text.replace(old, new)

    # Table 2: remove the withdrawn Liu et al. manuscript and update the count.
    literature = document.tables[1]
    delete_row(literature, 4)
    replace_all(document, "16 representative studies", "15 verified representative studies")
    replace_all(document, "16 representative papers", "15 verified representative papers")

    # Table 6: exact behavior after the code correction.
    architecture = document.tables[5]
    architecture.rows[7].cells[1].text = "Per-step delay, PDR, signal, and energy violations"
    architecture.rows[7].cells[2].text = "Persist and update multipliers across episode boundaries"
    architecture.rows[7].cells[3].text = "Training-wide adaptive penalty"

    # Algorithm 2: Stable-Baselines3 updates after the 200-step rollout, not
    # separately at every 50-step episode boundary.
    algorithm_2 = document.tables[8].cell(0, 0)
    algorithm_2.text = (
        "Algorithm 2. Adaptive Fuzzy LP-PPO training\n"
        "Input: scenarios Ω, episodes N_ep, horizon H, rollout length K=200, actor θ, critic φ, η_λ\n"
        "Output: trained policies {π_θ*}\n"
        "1: for each scenario ω ∈ Ω do\n"
        "2:    Initialize actor, critic and persistent λ = [0.35, 0.20, 0.60, 0.35]\n"
        "3:    Perform fuzzy-guided warm start using 800 samples and six epochs\n"
        "4:    Initialize the on-policy rollout buffer\n"
        "5:    for episode = 1, ..., N_ep do\n"
        "6:       Reset mobility, UAV positions and energy; retain λ\n"
        "7:       for t = 0, ..., H−1 do\n"
        "8:          Observe s_t and execute Algorithm 1\n"
        "9:          Sample a_t ~ π_θ(a_t | s_t) and map it to the effective UAV\n"
        "10:         Execute communication and calculate r_t\n"
        "11:         Compute violations and update persistent λ using Equation (30)\n"
        "12:         Update energy, vehicle position and next state\n"
        "13:         Store the transition in the on-policy rollout buffer\n"
        "14:         if K=200 rollout steps have been collected then\n"
        "15:             Compute GAE and update actor and critic for 10 PPO epochs\n"
        "16:             Clear the rollout buffer\n"
        "17:         end if\n"
        "18:      end for\n"
        "19:      Record episode metrics\n"
        "20:   end for\n"
        "21:   Save the scenario-specific policy\n"
        "22: end for"
    )

    # Remove experience-replay/global-synchronization claims. The implementation
    # trains three independent on-policy PPO models.
    replacements = {
        "Figure X : Working Architecture of the Proposed Fuzzy Lagrangian PPO Framework": "Figure 6. Working architecture of the proposed Fuzzy Lagrangian PPO framework.",
        "Figure Y : Intelligent Network Model of the Proposed FLP-PPO Framework": "Figure 7. Conceptual deployment topology of the proposed Fuzzy LP-PPO framework.",
    }
    for old, new in replacements.items():
        replace_all(document, old, new)

    set_text(
        document.paragraphs[182],
        "The interaction tuples (s_t, a_t, r_t, s_{t+1}) are stored temporarily in the on-policy PPO rollout buffer. Transitions are accumulated across episode boundaries until the configured 200-step rollout is complete; they are then used to compute GAE and perform ten PPO epochs before the buffer is cleared. This is an on-policy rollout mechanism rather than a replay memory."
    )
    set_text(
        document.paragraphs[183],
        "For each scenario, the centralized PPO actor–critic core contains one actor network, one critic network, GAE, and clipped PPO optimization. The actor estimates π_θ(a|s), the critic estimates V_φ(s), and the optimizer updates the scenario-specific model after each 200-step rollout. The urban-canyon, suburban-crossroads, and emergency-response policies are trained independently; no global cross-scenario parameter-synchronization bus is used."
    )
    set_text(
        document.paragraphs[181],
        "The Adaptive Lagrangian Constraint Bank updates persistent multipliers from observed QoS violations and incorporates the resulting penalties into the reward. This mechanism increases the cost of repeatedly violated constraints and encourages the policy to reduce delay, PDR, signal, and energy violations; it does not by itself guarantee zero violations under every network state."
    )
    set_text(
        document.paragraphs[194],
        "Figure 7 illustrates a conceptual deployment topology for the proposed centralized selector. It emphasizes vehicle-to-UAV candidate links, UAV coverage, communication measurements, and the edge-hosted decision pipeline. The topology is illustrative; the implemented simulator does not model distributed UAV agents or inter-UAV message passing."
    )
    set_text(
        document.paragraphs[196],
        "Multiple UAVs provide candidate coverage above each vehicular environment. Dashed links represent candidate vehicle-to-UAV associations, while a highlighted path represents the effective UAV chosen from the ranked list. UAV-to-UAV coordination and load-balancing protocols are outside the scope of the present centralized simulation."
    )
    set_text(
        document.paragraphs[202],
        "During training, interactions are accumulated in the on-policy rollout buffer. Once 200 steps have been collected, GAE is computed and mini-batches are reused for ten PPO optimization epochs; the buffer is then cleared before new on-policy data are collected."
    )
    set_text(
        document.paragraphs[203],
        "The resulting actor and critic parameters are stored as one trained model per scenario. At deployment, the appropriate scenario-specific actor is evaluated centrally and its discrete action is mapped to a position in the current ranked candidate list. No shared global controller or policy distribution to individual UAV agents is implemented."
    )
    set_text(
        document.paragraphs[200],
        "The persistent Lagrangian multipliers for delay, PDR, signal quality, and residual energy are updated whenever their corresponding violation magnitudes are positive. These adaptive penalties enter the reward used by PPO, encouraging lower violation severity over training. The clipped PPO update uses policy loss, value loss, entropy regularization, and GAE after every 200 collected steps."
    )

    # Remove the withdrawn paper from references and update verified records.
    for paragraph in list(document.paragraphs):
        if paragraph.text.startswith("[4] X. Liu"):
            delete_paragraph(paragraph)

    reference_updates = {
        "A. Uddin, A. H. Sakr, and N. Zhang, \"Task Offloading in Vehicular Edge Computing using Deep Reinforcement Learning: A Survey,\" arXiv:2502.06963, 2025.":
            "A. Uddin, A. H. Sakr, and N. Zhang, \"Intelligent Offloading in Vehicular Edge Computing: A Comprehensive Review of Deep Reinforcement Learning Approaches and Architectures,\" arXiv:2502.06963, 2025.",
        "A. Alagha, M. Kadadha, R. Mizouni, S. Singh, J. Bentahar, and H. Otrok, \"UAV-assisted Internet of Vehicles: A Framework Empowered by Reinforcement Learning and Blockchain,\" arXiv:2502.15713, 2025.":
            "A. Alagha, M. Kadadha, R. Mizouni, S. Singh, J. Bentahar, and H. Otrok, \"UAV-assisted Internet of Vehicles: A Framework Empowered by Reinforcement Learning and Blockchain,\" Vehicular Communications, 2025, doi:10.1016/j.vehcom.2025.100874.",
        "B. Li, W. Xie, Y. Ye, L. Liu, and Z. Fei, \"FlexEdge: Digital Twin-Enabled Task Offloading for UAV-Aided Vehicular Edge Computing,\" arXiv:2305.01536, 2023.":
            "B. Li, W. Xie, Y. Ye, L. Liu, and Z. Fei, \"FlexEdge: Digital Twin-Enabled Task Offloading for UAV-Aided Vehicular Edge Computing,\" IEEE Transactions on Vehicular Technology, 2023, doi:10.1109/TVT.2023.3262261.",
    }
    for paragraph in document.paragraphs:
        body = re.sub(r"^\[\d+\]\s*", "", paragraph.text)
        if body in reference_updates:
            set_text(paragraph, reference_updates[body])

    # Renumber the remaining bibliography contiguously after deleting unsupported
    # and withdrawn records. In-text citations use author-year format.
    references_heading_index = next(
        i for i, p in enumerate(document.paragraphs) if p.text.strip() == "References"
    )
    reference_paragraphs = [
        p for p in document.paragraphs[references_heading_index + 1 :] if p.text.strip()
    ]
    for number, paragraph in enumerate(reference_paragraphs, 1):
        body = re.sub(r"^\[\d+\]\s*", "", paragraph.text)
        set_text(paragraph, f"[{number}] {body}")

    document.core_properties.title = "Fuzzy LP-PPO for Reliable UAV Selection in Dynamic UAV-IoV Systems"
    document.core_properties.comments = "Implementation-faithful revision after active-constraint retraining"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
