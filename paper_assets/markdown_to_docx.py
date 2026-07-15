#!/usr/bin/env python3
"""Convert this paper's Markdown source to an editable Word document."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("`", "")
    text = re.sub(r"(?<!\\)\*([^*]+)\*", r"\1", text)
    return text.strip()


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(clean_inline(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(normalized):
        for c_idx, value in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))
    document.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw_rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        values = [part.strip() for part in lines[i].strip().strip("|").split("|")]
        raw_rows.append(values)
        i += 1
    rows = []
    for row in raw_rows:
        if all(re.fullmatch(r":?-{3,}:?", value or "") for value in row):
            continue
        rows.append(row)
    return rows, i


def add_body_paragraph(document: Document, text: str, style=None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.add_run(clean_inline(text))


def convert(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for name, size in (("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)):
        styles[name].font.name = "Times New Roman"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            i += 1
            continue

        image_match = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            image_path = (source.parent / image_match.group(2)).resolve()
            if image_path.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(image_path), width=Inches(6.2))
            else:
                add_body_paragraph(document, f"[Image unavailable: {image_match.group(1)}]")
            i += 1
            continue

        if stripped == "$$":
            equation = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                equation.append(lines[i].rstrip())
                i += 1
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("\n".join(equation).strip())
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(text)
            else:
                document.add_heading(text, level=min(level - 1, 3))
            i += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ordered:
            add_body_paragraph(document, ordered.group(1), style="List Number")
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            add_body_paragraph(document, bullet.group(1), style="List Bullet")
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt == "$$"
                or nxt.startswith("#")
                or nxt.startswith("|")
                or re.fullmatch(r"!\[([^]]*)\]\([^)]+\)", nxt)
                or re.match(r"^\d+\.\s+", nxt)
                or re.match(r"^[-*]\s+", nxt)
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        add_body_paragraph(document, " ".join(paragraph_lines))

    document.core_properties.title = "Fuzzy Lagrangian PPO for Reliable UAV Selection in IoV"
    document.core_properties.subject = "Research manuscript"
    document.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: markdown_to_docx.py SOURCE.md OUTPUT.docx")
    convert(Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve())
